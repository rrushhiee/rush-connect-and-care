from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
TMP_DIR = ROOT / "tmp" / "doc-assets"
SOURCE_LOGO = ROOT / "logo-512.png"
SMALL_LOGO = TMP_DIR / "rcc-logo-small.png"


def ensure_small_logo() -> Path:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    with Image.open(SOURCE_LOGO) as image:
        image = image.convert("RGBA")
        image.thumbnail((180, 180))
        image.save(SMALL_LOGO, optimize=True)
    return SMALL_LOGO


def add_logo_to_docx(path: Path, logo_path: Path) -> None:
    document = Document(path)

    for section in document.sections:
        header = section.header
        header.is_linked_to_previous = False

        for element in list(header._element):
            header._element.remove(element)

        paragraph = header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(str(logo_path), width=Inches(0.75))

        if section.top_margin < Inches(1):
            section.top_margin = Inches(1)

        section.header_distance = Inches(0.2)

    document.save(path)


def add_logo_to_xlsx(path: Path, logo_path: Path) -> None:
    workbook = load_workbook(path)

    for worksheet in workbook.worksheets:
        first_non_empty = None
        for row_index in range(1, min(worksheet.max_row, 12) + 1):
            if any(worksheet.cell(row_index, col).value not in (None, "") for col in range(1, worksheet.max_column + 1)):
                first_non_empty = row_index
                break

        if first_non_empty is None or first_non_empty <= 2:
            worksheet.insert_rows(1, amount=4)

        worksheet.row_dimensions[1].height = 24
        worksheet.row_dimensions[2].height = 24
        worksheet.row_dimensions[3].height = 24
        worksheet.row_dimensions[4].height = 12

        worksheet._images = []
        image = XLImage(str(logo_path))
        image.width = 72
        image.height = 72

        anchor_column = max(1, min(max(worksheet.max_column // 2, 2), 6))
        anchor = f"{get_column_letter(anchor_column)}1"
        worksheet.add_image(image, anchor)

    workbook.save(path)


def build_pdf_overlay(width: float, height: float, logo_path: Path) -> bytes:
    packet = io.BytesIO()
    pdf = canvas.Canvas(packet, pagesize=(width, height))

    logo_size = 54
    x_pos = (width - logo_size) / 2
    y_pos = height - logo_size - 10
    pdf.drawImage(str(logo_path), x_pos, y_pos, width=logo_size, height=logo_size, preserveAspectRatio=True, mask="auto")
    pdf.save()

    return packet.getvalue()


def add_logo_to_pdf(path: Path, logo_path: Path) -> None:
    reader = PdfReader(str(path))
    writer = PdfWriter()

    for page in reader.pages:
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        overlay_reader = PdfReader(io.BytesIO(build_pdf_overlay(width, height, logo_path)))
        overlay_page = overlay_reader.pages[0]
        page.merge_page(overlay_page)
        writer.add_page(page)

    with path.open("wb") as output_file:
        writer.write(output_file)


def main() -> None:
    logo_path = ensure_small_logo()

    for path in sorted(DOCS_DIR.iterdir()):
        if path.suffix == ".docx":
            add_logo_to_docx(path, logo_path)
        elif path.suffix == ".xlsx":
            add_logo_to_xlsx(path, logo_path)
        elif path.suffix == ".pdf":
            add_logo_to_pdf(path, logo_path)


if __name__ == "__main__":
    main()
